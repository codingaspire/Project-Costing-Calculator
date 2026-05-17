from flask import Flask, render_template, request, send_file
from docx import Document
from io import BytesIO

app = Flask(__name__)

def calculate_mandays(slot_of_designation, num_of_days):
    if slot_of_designation == 1:
        per_day_cost = 10500
    elif slot_of_designation == 2:
        per_day_cost = 6000
    elif slot_of_designation == 3:
        per_day_cost = 4500
    else:
        per_day_cost = 0

    mandays = per_day_cost * num_of_days
    return mandays

def calculate_hra(salary, hra_percentage):
    hra = (hra_percentage * salary) / 100
    return hra

def calculate_payment(salary, hra, months):
    payment = (salary + hra) * months
    return payment

def calculate_lab_share(direct_cost):
    lab_share = 0.875 * direct_cost
    return lab_share

def calculate_project_fee(direct_cost):
    project_fee = 0.625 * direct_cost
    return project_fee

def calculate_direct_cost(permanent_total, temporary_total, consumables, psu, equipment_cost, ta_da_national, ta_da_international, contingency, miscellaneous_cost_inputs):
    direct_cost = permanent_total + temporary_total + consumables + psu + equipment_cost + ta_da_national + ta_da_international + contingency + miscellaneous_cost_inputs
    return direct_cost

def calculate_project_cost(direct_cost, lab_share, project_fee, capital_cost, miscellaneous_cost):
    project_cost = direct_cost + lab_share + project_fee + capital_cost + miscellaneous_cost
    return project_cost

def calculate_total_tax(tax_percentage, project_cost):
    total_tax = (tax_percentage / 100) * project_cost
    return total_tax

def generate_word_document(permanent_staff, temporary_staff, cost_data):
    document = Document()
    
    document.add_heading('Cost Summary', level=1)
    
    table_1 = document.add_table(rows=2, cols=5)
    table_1.style = 'Table Grid'
    
    table1_header = table_1.rows[0].cells
    table1_header[0].text = 'a) Cost of Man-days'
    table1_header[1].text = 'i) Permanent Staff'
    table1_header[0].merge(table1_header[4])
    
    table_1.cell(1, 0).text = 'S.No.'
    table_1.cell(1, 1).text = 'Name & Designation'
    table_1.cell(1, 2).text = 'No. of days'
    table_1.cell(1, 3).text = 'Man-days Rate (Rs./day)'
    table_1.cell(1, 4).text = 'Amount (Rs)'
    
    for i, staff in enumerate(permanent_staff):
        row = table_1.add_row().cells
        row[0].text = str(i + 1)
        row[1].text = f"{staff['name']} (Slot {staff['slot_of_designation']})"
        row[2].text = str(staff['num_of_days'])
        if staff['slot_of_designation'] == 1:
            row[3].text = '10500'
        elif staff['slot_of_designation'] == 2:
            row[3].text = '6000'
        elif staff['slot_of_designation'] == 3:
            row[3].text = '4500'
        else:
            row[3].text = '0'
        row[4].text = str(staff['amount'])
    
    if temporary_staff:
        table_2 = document.add_table(rows=2, cols=7)
        table_2.style = 'Table Grid'
        
        table2_header = table_2.rows[0].cells
        table2_header[0].text = 'ii) Temporary Staff'
        table2_header[0].merge(table2_header[6])
        
        table_2.cell(1, 0).text = 'S.No.'
        table_2.cell(1, 1).text = 'Name'
        table_2.cell(1, 2).text = 'Post'
        table_2.cell(1, 3).text = 'Salary'
        table_2.cell(1, 4).text = 'HRA'
        table_2.cell(1, 5).text = 'Months'
        table_2.cell(1, 6).text = 'Payment'
        
        for i, staff in enumerate(temporary_staff):
            row = table_2.add_row().cells
            row[0].text = str(i + 1)
            row[1].text = staff['name']
            row[2].text = staff['post']
            row[3].text = str(staff['salary'])
            row[4].text = str(staff['hra'])
            row[5].text = str(staff['months'])
            row[6].text = str(staff['payment'])
    
    table = document.add_table(rows=9, cols=2)
    table.style = 'Table Grid'
    
    table.cell(0, 0).text = 'Parameter'
    table.cell(0, 1).text = 'Value'
    
    table.cell(1, 0).text = 'Direct Cost'
    table.cell(1, 1).text = str(cost_data['direct_cost'])
    
    table.cell(2, 0).text = 'Lab Share'
    table.cell(2, 1).text = str(cost_data['lab_share'])
    
    table.cell(3, 0).text = 'Project Fee'
    table.cell(3, 1).text = str(cost_data['project_fee'])
    
    table.cell(4, 0).text = 'Capital Cost'
    table.cell(4, 1).text = str(cost_data['capital_cost'])
    
    table.cell(5, 0).text = 'Miscellaneous Cost'
    table.cell(5, 1).text = str(cost_data['miscellaneous_cost'])
    
    table.cell(6, 0).text = 'Project Cost'
    table.cell(6, 1).text = str(cost_data['project_cost'])
    
    table.cell(7, 0).text = 'Total Tax'
    table.cell(7, 1).text = str(cost_data['total_tax'])
    
    table.cell(8, 0).text = 'Project Charges'
    table.cell(8, 1).text = str(cost_data['project_charges'])
    
    return document

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/calculate', methods=['POST'])
def calculate():
    permanent_staff = []
    permanent_count = int(request.form.get('permanent_count', 0))
    
    for i in range(permanent_count):
        name = request.form.get(f'permanent_name_{i}', '')
        slot = int(request.form.get(f'permanent_slot_{i}', 0))
        days = int(request.form.get(f'permanent_days_{i}', 0))
        amount = calculate_mandays(slot, days)
        permanent_staff.append({
            'name': name,
            'slot_of_designation': slot,
            'num_of_days': days,
            'amount': amount
        })
    
    temporary_staff = []
    temporary_count = int(request.form.get('temporary_count', 0))
    
    for i in range(temporary_count):
        name = request.form.get(f'temporary_name_{i}', '')
        post = request.form.get(f'temporary_post_{i}', '')
        salary = int(request.form.get(f'temporary_salary_{i}', 0))
        hra_given = int(request.form.get(f'temporary_hra_given_{i}', 0))
        if hra_given == 1:
            hra_percent = int(request.form.get(f'temporary_hra_percent_{i}', 0))
            hra = calculate_hra(salary, hra_percent)
        else:
            hra = 0
        months = int(request.form.get(f'temporary_months_{i}', 0))
        payment = calculate_payment(salary, hra, months)
        temporary_staff.append({
            'name': name,
            'post': post,
            'salary': salary,
            'hra': hra,
            'months': months,
            'payment': payment
        })
    
    consumables = int(request.form.get('consumables', 0))
    psu = int(request.form.get('psu', 0))
    equipment_cost = int(request.form.get('equipment_cost', 0))
    ta_da_national = int(request.form.get('ta_da_national', 0))
    ta_da_international = int(request.form.get('ta_da_international', 0))
    contingency = int(request.form.get('contingency', 0))
    miscellaneous_cost_inputs = int(request.form.get('miscellaneous_cost_inputs', 0))
    
    permanent_total = sum(staff['amount'] for staff in permanent_staff)
    temporary_total = sum(staff['payment'] for staff in temporary_staff)
    
    direct_cost = calculate_direct_cost(permanent_total, temporary_total, consumables, psu, equipment_cost, ta_da_national, ta_da_international, contingency, miscellaneous_cost_inputs)
    lab_share = calculate_lab_share(direct_cost)
    project_fee = calculate_project_fee(direct_cost)
    capital_cost = int(request.form.get('capital_cost', 0))
    miscellaneous_cost = int(request.form.get('miscellaneous_cost', 0))
    project_cost = calculate_project_cost(direct_cost, lab_share, project_fee, capital_cost, miscellaneous_cost)
    tax_percentage = int(request.form.get('tax_percentage', 0))
    total_tax = calculate_total_tax(tax_percentage, project_cost)
    project_charges = project_cost + total_tax
    
    cost_data = {
        'direct_cost': direct_cost,
        'lab_share': lab_share,
        'project_fee': project_fee,
        'capital_cost': capital_cost,
        'miscellaneous_cost': miscellaneous_cost,
        'project_cost': project_cost,
        'total_tax': total_tax,
        'project_charges': project_charges
    }
    
    return render_template('result.html', 
                          permanent_staff=permanent_staff,
                          temporary_staff=temporary_staff,
                          cost_data=cost_data)

@app.route('/download-docx', methods=['POST'])
def download_docx():
    import json
    permanent_staff_str = request.form.get('permanent_staff', '[]')
    temporary_staff_str = request.form.get('temporary_staff', '[]')
    cost_data_str = request.form.get('cost_data', '{}')
    
    permanent_staff = json.loads(permanent_staff_str)
    temporary_staff = json.loads(temporary_staff_str)
    cost_data = json.loads(cost_data_str)
    
    document = generate_word_document(permanent_staff, temporary_staff, cost_data)
    
    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    
    return send_file(file_stream, 
                    as_attachment=True, 
                    download_name='cost_summary.docx',
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

if __name__ == '__main__':
    app.run()
