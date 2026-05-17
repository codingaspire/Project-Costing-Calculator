from docx import Document

def get_valid_integer_input(prompt):
    while True:
        try:
            value = int(input(prompt))
            return value
        except ValueError:
            print("Invalid input. Please enter a valid integer.")

def calculate_mandays(slot_of_designation, num_of_days):
    if slot_of_designation == 1:
        per_day_cost = 10500
    elif slot_of_designation == 2:
        per_day_cost = 6000
    elif slot_of_designation == 3:
        per_day_cost = 4500
    else:
        per_day_cost = 0

    mandays = per_day_cost * num_of_days
    return mandays

def get_valid_alphabetic_input(prompt):
    while True:
        value = input(prompt)
        if value.isalpha():
            return value
        print("Invalid input. Please enter alphabetic characters only.")

def calculate_hra(salary, hra_percentage):
    hra = (hra_percentage * salary) / 100
    return hra

def calculate_payment(salary, hra, months):
    payment = (salary + hra) * months
    return payment

def calculate_lab_share(direct_cost):
    lab_share = 0.875 * direct_cost
    return lab_share

def calculate_project_fee(direct_cost):
    project_fee = 0.625 * direct_cost
    return project_fee

def calculate_direct_cost(amount_final, consumables, psu, equipment_cost, ta_da_national, ta_da_international, contingency, miscellaneous_cost_inputs):
    direct_cost = amount_final + consumables + psu + equipment_cost + ta_da_national + ta_da_international + contingency + miscellaneous_cost_inputs
    return direct_cost

def calculate_project_cost(direct_cost, lab_share, project_fee, capital_cost, miscellaneous_cost):
    project_cost = direct_cost + lab_share + project_fee + capital_cost + miscellaneous_cost
    return project_cost

def get_valid_tax_percentage(prompt):
    while True:
        try:
            tax_percentage = int(input(prompt))
            return tax_percentage
        except ValueError:
            print("Invalid input. Please enter a valid integer for the tax percentage.")

def calculate_total_tax(tax_percentage, project_cost):
    total_tax = (tax_percentage / 100) * project_cost
    return total_tax

num_staff = get_valid_integer_input("Enter the number of permanent staff: ")

staff_details = []
for i in range(num_staff):
    print(f"\nEnter details for Staff #{i+1}:")
    name = get_valid_alphabetic_input("Name: ")
    slot_of_designation = get_valid_integer_input("Slot of Designation (1, 2, 3): ")
    num_of_days = get_valid_integer_input("Number of days: ")

    amount = calculate_mandays(slot_of_designation, num_of_days)

    staff_details.append({
        "name": name,
        "slot_of_designation": slot_of_designation,
        "num_of_days": num_of_days,
        "amount": amount
    })

    print(f"Amount for {name}: {amount}")

# Access the staff details list for further processing if needed
print("\nStaff Details:")
for staff in staff_details:
    print(f"Name: {staff['name']}, Slot of Designation: {staff['slot_of_designation']}, Number of Days: {staff['num_of_days']}, Amount: {staff['amount']}")

num_staff = get_valid_integer_input("Enter the number of temporary staff: ")

staff_details = []
for i in range(num_staff):
    print(f"\nEnter details for Staff #{i+1}:")
    name = get_valid_alphabetic_input("Name: ")
    post = get_valid_alphabetic_input("Post: ")
    salary = get_valid_integer_input("Salary: ")

    hra_given = get_valid_integer_input("Does the staff receive HRA? (Press 1 for Yes, 2 for No): ")
    if hra_given == 1:
        hra_percentage = get_valid_integer_input("HRA percentage: ")
        hra = calculate_hra(salary, hra_percentage)
    else:
        hra = 0

    months = get_valid_integer_input("Number of months: ")
    payment = calculate_payment(salary, hra, months)

    staff_details.append({
        "name": name,
        "post": post,
        "salary": salary,
        "hra": hra,
        "months": months,
        "payment": payment
    })

    print(f"Salary for {name}: {salary}")

# Access the staff details list for further processing if needed
print("\nStaff Details:")
for staff in staff_details:
    print(f"Name: {staff['name']}, Post: {staff['post']}, Salary: {staff['salary']}, HRA: {staff['hra']}, Months: {staff['months']}, Payment: {staff['payment']}")

# Get initial cost inputs
consumables = get_valid_integer_input("Enter the cost of consumables: ")
psu = get_valid_integer_input("Enter the cost of Physical Inputs/Services/Utilities: ")
equipment_cost = get_valid_integer_input("Enter the equipment usage cost: ")
ta_da_national = get_valid_integer_input("Enter the TA/DA (national) cost: ")
ta_da_international = get_valid_integer_input("Enter the TA/DA (international) cost: ")
contingency = get_valid_integer_input("Enter the contingency cost: ")
miscellaneous_cost_inputs = get_valid_integer_input("Enter the cost of miscellaneous element (if any): ")


# Calculate direct cost
direct_cost = calculate_direct_cost(amount, consumables, psu, equipment_cost, ta_da_national, ta_da_international, contingency, miscellaneous_cost_inputs)

# Calculate lab share and project fee
project_fee = calculate_project_fee(direct_cost)
lab_share = calculate_lab_share(direct_cost)

# Get capital cost input
capital_cost = get_valid_integer_input("Enter the capital cost: ")

#direct_cost = calculate_direct_cost(consumables, psu, equipment_cost, ta_da_national, ta_da_international, contingency, miscellaneous_cost_inputs)
calculate_lab_share(direct_cost)
calculate_project_fee(direct_cost)
miscellaneous_cost = get_valid_integer_input("Enter the cost of miscellaneous element (if any): ")


project_cost = calculate_project_cost(direct_cost, lab_share, project_fee, capital_cost,miscellaneous_cost)

tax_percentage = get_valid_tax_percentage("Enter the tax percentage: ")
total_tax = calculate_total_tax(tax_percentage, project_cost)
project_charges = project_cost + total_tax

print("\nCost Summary:")
print(f"Direct Cost: {direct_cost}")
print(f"Lab Share: {lab_share}")
print(f"Project Fee: {project_fee}")
print(f"Capital Cost: {capital_cost}")
print(f"Miscellaneous Cost: {miscellaneous_cost}")
print(f"Project Cost: {project_cost}")
print(f"Total Tax: {total_tax}")
print(f"Project Charges: {project_charges}")

# Create a new Word document
document = Document()

# Add a title
document.add_heading('Cost Summary', level=1)

#Add a tabl_1
table_1 = document.add_table(rows=2, cols=5)
table_1.style = 'Table Grid'

# Table 1 Header Row
table1_header = table_1.rows[0].cells
table1_header[0].text = 'a) Cost of Man-days'
table1_header[1].text = 'i) Permanent Staff'
table1_header[0].merge(table1_header[4])

# Table 1 Column Headers
table_1.cell(1, 0).text = 'S.No.'
table_1.cell(1, 1).text = 'Name & Designation'
table_1.cell(1, 2).text = 'No. of days'
table_1.cell(1, 3).text = 'Man-days Rate (Rs./day)'
table_1.cell(1, 4).text = 'Amount (Rs)'

# Table 1 Data Row
for i, staff in enumerate(staff_details):
    row = table_1.add_row().cells
    row[0].text = str(i + 1)
    row[1].text = staff['name']
    row[2].text = str(staff['num_of_days'])

# Add a table
table = document.add_table(rows=9, cols=2)
table.style = 'Table Grid'

# Set the table headers
table.cell(0, 0).text = 'Parameter'
table.cell(0, 1).text = 'Value'

# Add the calculated values to the table
table.cell(1, 0).text = 'Direct Cost'
table.cell(1, 1).text = str(direct_cost)

table.cell(2, 0).text = 'Lab Share'
table.cell(2, 1).text = str(lab_share)

table.cell(3, 0).text = 'Project Fee'
table.cell(3, 1).text = str(project_fee)

table.cell(4, 0).text = 'Capital Cost'
table.cell(4, 1).text = str(capital_cost)

table.cell(5, 0).text = 'Miscellaneous Cost'
table.cell(5, 1).text = str(miscellaneous_cost)

table.cell(6, 0).text = 'Project Cost'
table.cell(6, 1).text = str(project_cost)

table.cell(7, 0).text = 'Total Tax'
table.cell(7, 1).text = str(total_tax)

table.cell(8, 0).text = 'Project Charges'
table.cell(8, 1).text = str(project_charges)

# Save the document
document.save('b_cost_summary.docx')